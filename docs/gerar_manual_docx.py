# -*- coding: utf-8 -*-
"""Gera o Manual do Usuário do sistema Escala em formato .docx."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRETO = RGBColor(0x1A, 0x1A, 0x1A)
CINZA = RGBColor(0x4B, 0x55, 0x63)


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_titulo_capa(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Manual do Usuário')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PRETO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Sistema Escala')
    run2.font.size = Pt(24)
    run2.font.color.rgb = CINZA

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Gestão de voluntários, ministérios, eventos e escalas de serviço')
    run3.font.size = Pt(13)
    run3.italic = True
    run3.font.color.rgb = CINZA

    for _ in range(10):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('Versão 1.0')
    run4.font.size = Pt(11)
    run4.font.color.rgb = CINZA

    doc.add_page_break()


def add_h1(doc, texto):
    h = doc.add_heading(level=1)
    run = h.add_run(texto)
    run.font.color.rgb = PRETO
    return h


def add_h2(doc, texto):
    h = doc.add_heading(level=2)
    run = h.add_run(texto)
    run.font.color.rgb = PRETO
    return h


def add_h3(doc, texto):
    h = doc.add_heading(level=3)
    run = h.add_run(texto)
    run.font.color.rgb = PRETO
    return h


def add_paragrafo(doc, texto, negrito=False, italico=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    return p


def add_bullet(doc, texto, negrito_prefixo=None):
    p = doc.add_paragraph(style='List Bullet')
    if negrito_prefixo:
        r1 = p.add_run(negrito_prefixo)
        r1.bold = True
        p.add_run(texto)
    else:
        p.add_run(texto)
    return p


def add_passo(doc, numero, texto):
    p = doc.add_paragraph(style='List Number')
    p.add_run(texto)
    return p


def add_tabela(doc, cabecalhos, linhas, larguras=None):
    tabela = doc.add_table(rows=1, cols=len(cabecalhos))
    tabela.style = 'Table Grid'
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tabela.rows[0].cells
    for i, titulo in enumerate(cabecalhos):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(titulo)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], '1A1A1A')
    for linha in linhas:
        row_cells = tabela.add_row().cells
        for i, valor in enumerate(linha):
            row_cells[i].text = str(valor)
    if larguras:
        for row in tabela.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tabela


def add_nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run('Nota: ')
    r.bold = True
    r.italic = True
    r2 = p.add_run(texto)
    r2.italic = True


def build():
    doc = Document()

    estilo_normal = doc.styles['Normal']
    estilo_normal.font.name = 'Calibri'
    estilo_normal.font.size = Pt(11)

    add_titulo_capa(doc)

    # Sumário simples (manual, sem TOC de campo)
    add_h1(doc, 'Sumário')
    itens_sumario = [
        '1. Introdução',
        '2. Papéis de usuário: Voluntário e Administrador',
        '3. Acessando o sistema',
        '   3.1 Login',
        '   3.2 Cadastro de novo voluntário',
        '4. Menu de navegação',
        '5. Área do Voluntário',
        '   5.1 Minha Escala',
        '   5.2 Indisponibilidade',
        '6. Área do Administrador',
        '   6.1 Ministérios',
        '   6.2 Voluntários',
        '   6.3 Eventos',
        '   6.4 Escala (geração e gestão)',
        '7. Notificações via WhatsApp',
        '8. Perguntas frequentes',
    ]
    for item in itens_sumario:
        add_paragrafo(doc, item)
    doc.add_page_break()

    # 1. Introdução
    add_h1(doc, '1. Introdução')
    add_paragrafo(
        doc,
        'O sistema Escala foi desenvolvido para organizar a escala de voluntários de uma igreja ou '
        'organização, facilitando o cadastro de ministérios, eventos (cultos/reuniões), a geração '
        'automática de escalas e a comunicação com os voluntários via WhatsApp.'
    )
    add_paragrafo(
        doc,
        'Este manual descreve todas as telas e funcionalidades disponíveis, tanto para o Voluntário '
        'quanto para o Administrador do sistema.'
    )

    # 2. Papéis
    add_h1(doc, '2. Papéis de usuário: Voluntário e Administrador')
    add_paragrafo(doc, 'O sistema possui dois tipos de perfil de acesso:')
    add_bullet(doc, ': pode visualizar e responder às suas próprias escalas e cadastrar suas datas de indisponibilidade.', 'Voluntário')
    add_bullet(doc, ': além de tudo que o voluntário pode fazer, gerencia ministérios, voluntários, eventos e a geração/edição das escalas de todos os voluntários.', 'Administrador')
    add_tabela(
        doc,
        ['Aspecto', 'Voluntário', 'Administrador'],
        [
            ['Tela inicial ao entrar', 'Minha Escala', 'Escala (admin)'],
            ['Itens do menu', 'Minha Escala, Indisponibilidade', 'Escala, Eventos, Voluntários, Ministérios'],
            ['Pode se cadastrar sozinho', 'Sim, pela tela de Registro', 'Não; é criado por um Administrador'],
            ['Gerencia ministérios/eventos/usuários', 'Não', 'Sim'],
            ['Gera escalas automaticamente', 'Não', 'Sim'],
            ['Responde à própria escala', 'Sim (Confirmar / Não posso)', 'Pode confirmar/recusar em nome de qualquer voluntário'],
        ],
    )

    # 3. Acessando o sistema
    add_h1(doc, '3. Acessando o sistema')

    add_h2(doc, '3.1 Login')
    add_paragrafo(doc, 'Tela inicial do sistema, disponível em /login.')
    add_paragrafo(doc, 'Campos do formulário:')
    add_bullet(doc, ': endereço de e-mail cadastrado.', 'E-mail')
    add_bullet(doc, ': senha de acesso.', 'Senha')
    add_paragrafo(doc, 'Passo a passo:')
    add_passo(doc, 1, 'Informe o e-mail e a senha cadastrados.')
    add_passo(doc, 2, 'Clique em "Entrar".')
    add_passo(doc, 3, 'Se as credenciais estiverem corretas, você será redirecionado automaticamente: administradores vão para a tela de Escala; voluntários vão para a tela Minha Escala.')
    add_nota(doc, 'se e-mail ou senha estiverem incorretos, uma mensagem de erro será exibida abaixo do formulário.')
    add_paragrafo(doc, 'Ainda não tem uma conta? Clique no link "Cadastre-se como voluntário" para ir à tela de Registro.')

    add_h2(doc, '3.2 Cadastro de novo voluntário')
    add_paragrafo(doc, 'Disponível em /registrar, para que novos voluntários criem sua própria conta.')
    add_paragrafo(doc, 'Campos do formulário:')
    add_bullet(doc, ': nome completo do voluntário.', 'Nome completo')
    add_bullet(doc, ': e-mail que será usado para login.', 'E-mail')
    add_bullet(doc, ': número de WhatsApp com código do país e DDD (ex.: 5511999999999), usado para receber notificações de escala.', 'Telefone/WhatsApp')
    add_bullet(doc, ': senha de acesso ao sistema.', 'Senha')
    add_bullet(doc, ': lista de ministérios em que o voluntário deseja servir (seleção múltipla por caixas de marcação).', 'Ministérios')
    add_paragrafo(doc, 'Passo a passo:')
    add_passo(doc, 1, 'Preencha todos os campos obrigatórios.')
    add_passo(doc, 2, 'Marque um ou mais ministérios de interesse.')
    add_passo(doc, 3, 'Clique em "Cadastrar".')
    add_passo(doc, 4, 'O sistema realiza o login automaticamente e leva você à tela Minha Escala.')
    add_nota(doc, 'todo cadastro feito por este formulário cria um usuário do tipo Voluntário. Contas de Administrador são criadas pela tela de gestão de Voluntários por outro administrador.')

    # 4. Menu de navegação
    add_h1(doc, '4. Menu de navegação')
    add_paragrafo(
        doc,
        'Após o login, uma barra de navegação (Navbar) fica visível no topo de todas as telas, exibindo '
        'o nome do usuário logado e o botão "Sair" (que encerra a sessão e retorna à tela de Login).'
    )
    add_paragrafo(doc, 'Em telas pequenas (celular), os links de navegação ficam ocultos em um menu do tipo hambúrguer (ícone ☰), que pode ser aberto/fechado tocando no ícone.')
    add_paragrafo(doc, 'Itens de menu por perfil:')
    add_bullet(doc, ': Minha Escala, Indisponibilidade.', 'Voluntário')
    add_bullet(doc, ': Escala, Eventos, Voluntários, Ministérios.', 'Administrador')

    # 5. Área do voluntário
    add_h1(doc, '5. Área do Voluntário')

    add_h2(doc, '5.1 Minha Escala')
    add_paragrafo(doc, 'Tela inicial do voluntário (rota /). Mostra a lista de escalas (atribuições) futuras do usuário logado.')
    add_paragrafo(doc, 'Colunas da tabela: Evento, Data, Ministério, Status e Ações.')
    add_paragrafo(doc, 'Status possíveis, exibidos como selos coloridos:')
    add_bullet(doc, ': o voluntário ainda não respondeu.', 'Pendente')
    add_bullet(doc, ': o voluntário confirmou presença.', 'Confirmado')
    add_bullet(doc, ': o voluntário informou que não poderá servir.', 'Recusado')
    add_paragrafo(doc, 'Para escalas com status "Pendente", dois botões ficam disponíveis:')
    add_bullet(doc, ': confirma a presença naquela escala.', 'Confirmar')
    add_bullet(doc, ': informa que não poderá comparecer; a vaga poderá ser reatribuída pelo administrador.', 'Não posso')
    add_paragrafo(doc, 'Se não houver escalas futuras, é exibida a mensagem "Você ainda não foi escalado para nenhum evento futuro."')

    add_h2(doc, '5.2 Indisponibilidade')
    add_paragrafo(doc, 'Disponível em /indisponibilidade. Permite que o voluntário informe datas em que não poderá servir, para que o sistema não o escale automaticamente nesses dias.')
    add_paragrafo(doc, 'Formulário de cadastro:')
    add_bullet(doc, ': data em que o voluntário estará indisponível.', 'Data')
    add_bullet(doc, ': texto livre e opcional (ex.: "Viagem", "Saúde").', 'Motivo (opcional)')
    add_passo(doc, 1, 'Selecione a data desejada.')
    add_passo(doc, 2, 'Preencha o motivo, se desejar.')
    add_passo(doc, 3, 'Clique em "Adicionar".')
    add_paragrafo(doc, 'A lista abaixo do formulário exibe todas as indisponibilidades já cadastradas, com um botão "Remover" para excluir qualquer registro.')

    # 6. Área do administrador
    add_h1(doc, '6. Área do Administrador')

    add_h2(doc, '6.1 Ministérios')
    add_paragrafo(doc, 'Disponível em /admin/ministerios. Permite cadastrar e remover os ministérios (equipes de serviço) da organização, como "Louvor", "Recepção", "Mídia", etc.')
    add_paragrafo(doc, 'Para cadastrar um novo ministério:')
    add_passo(doc, 1, 'Digite o nome do ministério no campo "Novo ministério".')
    add_passo(doc, 2, 'Clique em "Adicionar".')
    add_paragrafo(doc, 'Para remover um ministério, clique em "Remover" na linha correspondente e confirme a exclusão na janela de confirmação.')
    add_nota(doc, 'remover um ministério pode afetar eventos e escalas que dependem dele. Use com cautela.')

    add_h2(doc, '6.2 Voluntários')
    add_paragrafo(doc, 'Disponível em /admin/voluntarios. Permite cadastrar novos voluntários (inclusive por indicação do administrador) e gerenciar os já existentes.')
    add_h3(doc, 'Cadastro de novo voluntário (pelo administrador)')
    add_paragrafo(doc, 'Campos: Nome completo, E-mail, Telefone/WhatsApp, Senha provisória e seleção de Ministérios.')
    add_passo(doc, 1, 'Preencha os dados do voluntário.')
    add_passo(doc, 2, 'Marque os ministérios em que ele/ela poderá servir.')
    add_passo(doc, 3, 'Clique em "Cadastrar voluntário".')
    add_h3(doc, 'Gestão da lista de voluntários')
    add_paragrafo(doc, 'A tabela exibe: Nome, Contato (e-mail e telefone), Função (Admin/Voluntário), Ministérios, Status (Ativo/Inativo) e Ações.')
    add_bullet(doc, ': abre um modo de edição em que é possível marcar/desmarcar os ministérios daquele voluntário; use "Salvar" para confirmar ou "Cancelar" para descartar.', 'Editar ministérios')
    add_bullet(doc, ': alterna a situação do voluntário. Um voluntário inativo não é considerado nas gerações automáticas de escala.', 'Ativar / Desativar')
    add_bullet(doc, ': exclui o voluntário do sistema, após confirmação.', 'Remover')

    add_h2(doc, '6.3 Eventos')
    add_paragrafo(doc, 'Disponível em /admin/eventos. Permite cadastrar cultos, reuniões ou outros eventos que precisam de voluntários escalados, definindo quantas vagas cada ministério precisa preencher.')
    add_h3(doc, 'Criar um novo evento')
    add_passo(doc, 1, 'Informe o nome do evento (ex.: "Culto de Domingo").')
    add_passo(doc, 2, 'Informe a data e o horário do evento.')
    add_passo(doc, 3, 'Em "Necessidades", selecione um ministério e informe a quantidade de vagas necessárias.')
    add_passo(doc, 4, 'Clique em "+ Adicionar ministério" para incluir necessidades de outros ministérios, se necessário.')
    add_passo(doc, 5, 'Use o botão "Remover" ao lado de uma necessidade para excluí-la do formulário, se adicionada por engano.')
    add_passo(doc, 6, 'Clique em "Criar evento" para salvar.')
    add_h3(doc, 'Lista de eventos cadastrados')
    add_paragrafo(doc, 'Exibe nome, data e as necessidades de cada evento (em forma de etiquetas, ex.: "Louvor x3"). O botão "Remover" exclui o evento e todas as escalas associadas a ele, após confirmação.')

    add_h2(doc, '6.4 Escala (geração e gestão)')
    add_paragrafo(doc, 'Disponível em /admin/escala — tela principal do administrador, onde as escalas de voluntários são geradas e ajustadas.')
    add_h3(doc, 'Geração automática para os próximos 30 dias')
    add_paragrafo(
        doc,
        'O botão "Gerar escala automática (próximos 30 dias)" analisa todos os eventos cadastrados no período, '
        'distribui os voluntários entre as vagas necessárias de cada ministério (respeitando as indisponibilidades '
        'cadastradas) e envia automaticamente uma notificação via WhatsApp a cada voluntário escalado.'
    )
    add_h3(doc, 'Geração por evento específico')
    add_paragrafo(doc, 'Em cada card de evento, o botão "Gerar/completar escala deste evento" preenche apenas as vagas em falta daquele evento específico.')
    add_h3(doc, 'Tabela de escalas do evento')
    add_paragrafo(doc, 'Colunas: Ministério, Voluntário, Status, Notificado e Ações.')
    add_bullet(doc, ': permite reatribuir a vaga a outro voluntário elegível (que participe do mesmo ministério). A troca é aplicada imediatamente.', 'Campo Voluntário (lista suspensa)')
    add_bullet(doc, ': indica se a notificação via WhatsApp foi enviada com sucesso; passe o mouse sobre o selo para ver a data/hora do envio ou o motivo de uma eventual falha.', 'Selo "Notificado"')
    add_bullet(doc, ': define manualmente o status da escala como Confirmado, independentemente da resposta do voluntário.', 'Botão "Confirmar"')
    add_bullet(doc, ': define manualmente o status da escala como Recusado.', 'Botão "Recusar"')
    add_bullet(doc, ': envia novamente a mensagem de notificação ao voluntário (útil em caso de falha no primeiro envio).', 'Botão "Reenviar notificação"')
    add_bullet(doc, ': exclui aquela atribuição específica, após confirmação, liberando a vaga.', 'Botão "Remover"')

    # 7. WhatsApp
    add_h1(doc, '7. Notificações via WhatsApp')
    add_paragrafo(
        doc,
        'Sempre que uma escala é gerada, o voluntário responsável recebe uma mensagem via WhatsApp informando '
        'o ministério, o evento e a data, com a orientação de responder "SIM" para confirmar ou "NAO" para recusar '
        'diretamente pelo WhatsApp.'
    )
    add_paragrafo(
        doc,
        'Para evitar o bloqueio do número por envio em massa, o sistema envia as mensagens de forma sequencial, '
        'com um pequeno intervalo aleatório entre cada envio, especialmente importante quando muitos voluntários '
        'são notificados de uma só vez (por exemplo, ao usar a geração automática para os próximos 30 dias).'
    )
    add_paragrafo(
        doc,
        'O status de envio de cada notificação pode ser consultado na coluna "Notificado" da tela de Escala do '
        'Administrador, incluindo a data/hora do envio ou o motivo de uma falha.'
    )

    # 8. FAQ
    add_h1(doc, '8. Perguntas frequentes')
    perguntas = [
        ('Não recebi a notificação da minha escala pelo WhatsApp, o que fazer?',
         'Verifique se seu número de telefone está correto no seu cadastro. Solicite ao administrador que utilize o botão "Reenviar notificação" na tela de Escala.'),
        ('Como faço para não ser escalado em uma data específica?',
         'Acesse a tela "Indisponibilidade" e cadastre a data (e, se desejar, o motivo) antes que a escala seja gerada para aquele período.'),
        ('Já confirmei uma escala, mas não poderei mais comparecer. O que faço?',
         'Entre em contato com o administrador responsável para que ele ajuste manualmente o status ou reatribua a vaga a outro voluntário na tela de Escala.'),
        ('Esqueci minha senha, como recuperar o acesso?',
         'Atualmente o sistema não possui recuperação automática de senha pela tela de login; solicite a um administrador que atualize seus dados de acesso.'),
        ('Posso participar de mais de um ministério?',
         'Sim. Tanto no cadastro (Registro) quanto na edição feita por um administrador, é possível selecionar múltiplos ministérios para o mesmo voluntário.'),
    ]
    for pergunta, resposta in perguntas:
        add_h3(doc, pergunta)
        add_paragrafo(doc, resposta)

    doc.add_paragraph()
    rodape = add_paragrafo(doc, 'Fim do manual — Sistema Escala.', italico=True)
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caminho = 'docs/Manual_do_Usuario_Escala.docx'
    doc.save(caminho)
    print('Gerado:', caminho)


if __name__ == '__main__':
    build()
